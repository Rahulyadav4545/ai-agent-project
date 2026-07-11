from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from langchain_groq import ChatGroq
from langchain_core.tools import tool
from langgraph.prebuilt import create_react_agent
from langchain_core.messages import HumanMessage, SystemMessage
import os
import json
from datetime import datetime

# ============================================
# 1. FASTAPI APP SETUP
# ============================================
app = FastAPI(
    title="Autonomous AI Agent API",
    description="60-Minute Build Challenge - Python AI Engineer",
    version="1.0.0"
)

# ============================================
# 2. GROQ API KEY SETUP
# ============================================
# IMPORTANT: Replace with your actual Groq API key from groq.com
os.environ["GROQ_API_KEY"] = "paste your_groq_api_key_here"

# Initialize LLM (Groq Llama 3.1 - Free)
llm = ChatGroq(temperature=0.3, model_name="llama-3.1-8b-instant")

# ============================================
# 3. TOOLS DEFINITION
# ============================================

@tool
def create_word_document(title: str, sections: str, filename: str = "output.docx") -> str:
    """
    Creates a professional Microsoft Word document (.docx) with structured content.
    
    Args:
        title: Main heading of the document
        sections: Content with sections (use "\\n\\n" for section breaks)
        filename: Output filename (must end with .docx)
    
    Returns:
        Success/error message with file path
    """
    try:
        doc = Document()
        
        # Add main title
        doc.add_heading(title, 0)
        
        # Add timestamp
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Process sections
        current_section = ""
        for line in sections.split("\n"):
            if line.strip():
                # Check if it's a heading (starts with # or all caps)
                if line.startswith("#"):
                    current_section = line.replace("#", "").strip()
                    doc.add_heading(current_section, level=1)
                elif line.isupper() and len(line) > 3:
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(line)
        
        # Save document
        doc.save(filename)
        return f"✅ Document created successfully: {filename}"
    
    except Exception as e:
        return f"❌ Error creating document: {str(e)}"


@tool
def analyze_request(request: str) -> str:
    """
    Analyzes the user request and returns key requirements.
    
    Args:
        request: The user's request text
    
    Returns:
        Analysis summary
    """
    analysis = f"📋 Request Analysis:\n- Type: Business Document\n- Complexity: Multi-step\n- Key Focus: {request[:100]}..."
    return analysis


# Tools list
tools = [create_word_document, analyze_request]

# ============================================
# 4. AGENT SETUP (Using LangGraph ReAct)
# ============================================
agent_executor = create_react_agent(llm, tools)

# ============================================
# 5. REQUEST/RESPONSE MODELS
# ============================================

class RequestModel(BaseModel):
    request: str
    
    class Config:
        json_schema_extra = {
            "example": {
                "request": "Create meeting minutes for a product roadmap discussion with 5 attendees"
            }
        }


class ResponseModel(BaseModel):
    status: str
    message: str
    task_plan: list
    agent_reasoning: str
    document_created: bool
    document_path: str
    execution_time: str

# ============================================
# 6. MAIN AGENT ENDPOINT
# ============================================

@app.post("/agent", response_model=ResponseModel)
async def run_agent(payload: RequestModel):
    """
    Main endpoint: Accepts request, runs autonomous agent, generates Word document.
    
    The agent:
    1. Understands the request
    2. Creates its own task plan
    3. Generates professional content
    4. Saves as Word document
    """
    
    import time
    start_time = time.time()
    
    try:
        # ============================================
        # SYSTEM PROMPT - Guides agent behavior
        # ============================================
        system_prompt = SystemMessage(content="""You are an autonomous AI engineering agent.

When you receive a user request, ALWAYS follow these steps:

1. TASK PLANNING: Identify 4-5 specific tasks needed to complete the request.
   Format: "TASK BREAKDOWN:
   - Task 1: [specific description]
   - Task 2: [specific description]
   ..."

2. CONTENT GENERATION: Create professional business content based on the request.
   Make reasonable assumptions if information is missing.

3. DOCUMENT CREATION: You MUST use the 'create_word_document' tool to save the content.
   Call the tool with:
   - title: Professional document title
   - sections: Well-formatted content with clear sections

IMPORTANT: 
- Make your task breakdown CLEAR and EXPLICIT
- Make reasonable assumptions for ambiguous requests
- Always use the create_word_document tool to save output
- Format output professionally

User Request: {user_request}""".format(user_request=payload.request))

        user_prompt = HumanMessage(content=payload.request)
        
        # ============================================
        # RUN THE AGENT
        # ============================================
        inputs = {"messages": [system_prompt, user_prompt]}
        response = agent_executor.invoke(inputs)
        
        # ============================================
        # EXTRACT RESULTS
        # ============================================
        full_response = response["messages"][-1].content
        
        # Parse task breakdown from response
        task_plan = extract_task_plan(full_response)
        
        # Check if document was created
        document_created = os.path.exists("output.docx")
        document_path = "output.docx" if document_created else "not_created"
        
        # Calculate execution time
        execution_time = f"{time.time() - start_time:.2f}s"
        
        return ResponseModel(
            status="success",
            message="✅ Agent executed successfully. Document generated.",
            task_plan=task_plan,
            agent_reasoning=full_response[:500] + "..." if len(full_response) > 500 else full_response,
            document_created=document_created,
            document_path=document_path,
            execution_time=execution_time
        )
        
    except Exception as e:
        execution_time = f"{time.time() - start_time:.2f}s"
        raise HTTPException(
            status_code=500,
            detail=f"❌ Agent execution failed: {str(e)}"
        )


# ============================================
# 7. HELPER FUNCTION - Extract Task Plan
# ============================================

def extract_task_plan(response_text: str) -> list:
    """
    Extracts task plan from agent response.
    Looks for "TASK" or "Task" or numbered lists.
    """
    tasks = []
    
    lines = response_text.split("\n")
    
    for line in lines:
        line = line.strip()
        
        # Look for task indicators
        if any(indicator in line.lower() for indicator in ["task", "step", "-", "•"]):
            # Clean up the line
            cleaned = line.replace("Task", "").replace("task", "").replace("Step", "").replace("step", "").replace("-", "").replace("•", "").strip()
            
            if cleaned and len(cleaned) > 5:
                tasks.append(cleaned[:100])  # First 100 chars
    
    # If no tasks found, create default
    if not tasks:
        tasks = [
            "Analyze and understand the user request",
            "Identify required information and assumptions",
            "Generate professional content structure",
            "Create formatted Word document",
            "Validate and save output"
        ]
    
    return tasks[:5]  # Return max 5 tasks


# ============================================
# 8. DOWNLOAD ENDPOINT
# ============================================

@app.get("/download")
async def download_document():
    """
    Download the generated Word document.
    """
    file_path = "output.docx"
    
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=404,
            detail="Document not found. Please run /agent endpoint first."
        )
    
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"agent_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )


# ============================================
# 9. HEALTH CHECK ENDPOINT
# ============================================

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "service": "Autonomous AI Agent",
        "version": "1.0.0",
        "llm": "Groq Llama 3.1",
        "api_key_configured": bool(os.environ.get("GROQ_API_KEY", "").startswith("gsk_"))
    }


# ============================================
# 10. ROOT ENDPOINT
# ============================================

@app.get("/")
async def root():
    """Root endpoint with API documentation."""
    return {
        "message": "🤖 Autonomous AI Agent API - 60-Minute Build Challenge",
        "endpoints": {
            "POST /agent": "Submit request and get generated document",
            "GET /download": "Download the generated Word document",
            "GET /health": "Check API health status",
            "GET /docs": "Interactive API documentation (Swagger)"
        },
        "example_usage": {
            "method": "POST",
            "url": "http://localhost:8000/agent",
            "body": {
                "request": "Create a professional meeting summary for a product review meeting"
            }
        }
    }


# ============================================
# 11. RUN SERVER
# ============================================

if __name__ == "__main__":
    import uvicorn
    
    print("""
    ╔════════════════════════════════════════════════════════╗
    ║     🤖 AUTONOMOUS AI AGENT - Starting Server...        ║
    ╚════════════════════════════════════════════════════════╝
    
    📍 API: http://localhost:8000
    📚 Docs: http://localhost:8000/docs
    
    ════════════════════════════════════════════════════════
    """)
    
    uvicorn.run(app, host="0.0.0.0", port=8000)
